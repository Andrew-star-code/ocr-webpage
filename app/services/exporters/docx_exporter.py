import io
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.schemas.recognition import HeadingBlock,ListBlock,TableBlock
_ALIGN={"left":WD_ALIGN_PARAGRAPH.LEFT,"center":WD_ALIGN_PARAGRAPH.CENTER,"right":WD_ALIGN_PARAGRAPH.RIGHT,"justify":WD_ALIGN_PARAGRAPH.JUSTIFY}
class DocxExporter:
 format="docx";mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document";extension="docx"
 def export(self,result,options):
  output=io.BytesIO();word=Document()
  for pi,page in enumerate(result.document.pages):
   if pi:word.add_page_break()
   for b in sorted(page.blocks,key=lambda x:x.reading_order):
    value=b.normalized_text or b.original_text
    if isinstance(b,HeadingBlock) and value:p=word.add_heading(value,level=b.heading_level)
    elif isinstance(b,ListBlock):
     for item in b.items:
      p=word.add_paragraph(style="List Number" if b.ordered else "List Bullet");p.paragraph_format.left_indent=None;p.add_run(item.text)
     continue
    elif isinstance(b,TableBlock):
     table=word.add_table(rows=b.row_count,cols=b.column_count);table.style="Table Grid"
     occupied=set()
     for row in b.rows:
      for c in row.cells:
       if (c.row_index,c.column_index) in occupied:continue
       cell=table.cell(c.row_index,c.column_index);cell.text=c.text
       end=(c.row_index+c.row_span-1,c.column_index+c.column_span-1)
       if end!=(c.row_index,c.column_index):cell.merge(table.cell(*end))
       occupied.update((r,col) for r in range(c.row_index,c.row_index+c.row_span) for col in range(c.column_index,c.column_index+c.column_span))
       if c.is_header:
        for run in cell.paragraphs[0].runs:run.bold=True
     continue
    elif b.type in {"checkbox","signature","stamp"}:
     label={"checkbox":"[Флажок]","signature":"[Подпись]","stamp":"[Печать]"}[b.type];p=word.add_paragraph(label+(f" {value}" if value else ""))
    elif value:p=word.add_paragraph(value)
    else:continue
    if options.preserve_layout:
     p.alignment=_ALIGN[b.style.alignment]
     for run in p.runs:run.bold=b.style.bold;run.italic=b.style.italic;run.underline=b.style.underline
  word.save(output);return output.getvalue()
