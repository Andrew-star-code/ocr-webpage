import html,io,json
import fitz
from docx import Document as DocxDocument
from app.schemas.recognition import DocumentRecognition,HeadingBlock,ListBlock,TableBlock

def text(doc):
    chunks=[]
    for page in doc.pages:
        chunks.append(f"--- Страница {page.page_number} ---")
        for b in sorted(page.blocks,key=lambda x:x.reading_order):
            if isinstance(b,ListBlock): chunks.extend((f"{i+1}. " if b.ordered else "• ")+x.text for i,x in enumerate(b.items))
            elif isinstance(b,TableBlock): chunks.extend(" | ".join(c.text for c in r.cells) for r in b.rows)
            elif b.original_text: chunks.append(b.normalized_text or b.original_text)
    return "\n\n".join(chunks).encode()
def markdown(doc):
    out=[]
    for page in doc.pages:
        out.append(f"<!-- page {page.page_number} -->\n\n---")
        for b in sorted(page.blocks,key=lambda x:x.reading_order):
            value=b.normalized_text or b.original_text
            if isinstance(b,HeadingBlock): out.append("#"*b.heading_level+" "+value)
            elif isinstance(b,ListBlock): out.extend((f"{i+1}. " if b.ordered else "- ")+x.text for i,x in enumerate(b.items))
            elif isinstance(b,TableBlock):
                rows=[[c.text for c in r.cells] for r in b.rows]
                if rows: out.extend(["| "+" | ".join(rows[0])+" |","| "+" | ".join("---" for _ in rows[0])+" |"]+["| "+" | ".join(r)+" |" for r in rows[1:]])
            elif value: out.append(value)
    return "\n\n".join(out).encode()
def html_export(doc):
    out=["<!doctype html><html><meta charset=utf-8><body>"]
    for page in doc.pages:
        out.append(f'<section data-page="{page.page_number}">')
        for b in sorted(page.blocks,key=lambda x:x.reading_order):
            value=html.escape(b.normalized_text or b.original_text)
            if isinstance(b,HeadingBlock): out.append(f"<h{b.heading_level}>{value}</h{b.heading_level}>")
            elif isinstance(b,ListBlock): out.append(("<ol>" if b.ordered else "<ul>")+"".join(f"<li>{html.escape(x.text)}</li>" for x in b.items)+("</ol>" if b.ordered else "</ul>"))
            elif isinstance(b,TableBlock): out.append("<table>"+"".join("<tr>"+"".join(("<th>" if c.is_header else "<td>")+html.escape(c.text)+("</th>" if c.is_header else "</td>") for c in r.cells)+"</tr>" for r in b.rows)+"</table>")
            elif value: out.append(f"<p>{value}</p>")
        out.append("</section>")
    return "".join(out+["</body></html>"]).encode()
def docx(doc):
    output=io.BytesIO(); word=DocxDocument()
    for pi,page in enumerate(doc.pages):
        if pi: word.add_page_break()
        for b in sorted(page.blocks,key=lambda x:x.reading_order):
            value=b.normalized_text or b.original_text
            if isinstance(b,HeadingBlock): p=word.add_heading(value,level=b.heading_level)
            elif isinstance(b,ListBlock):
                for item in b.items: word.add_paragraph(item.text,style="List Number" if b.ordered else "List Bullet")
                continue
            elif isinstance(b,TableBlock):
                table=word.add_table(rows=b.row_count,cols=b.column_count)
                for row in b.rows:
                    for c in row.cells:
                        cell=table.cell(c.row_index,c.column_index); cell.text=c.text
                        if c.row_span>1 or c.column_span>1: cell.merge(table.cell(c.row_index+c.row_span-1,c.column_index+c.column_span-1))
                continue
            else: p=word.add_paragraph(value)
            for run in p.runs: run.bold=b.style.bold; run.italic=b.style.italic; run.underline=b.style.underline
    word.save(output); return output.getvalue()
def json_export(doc): return doc.model_dump_json(indent=2).encode()
def searchable_pdf(doc,images):
    output=fitz.open()
    for page_data,image in zip(doc.pages,images):
        pix=fitz.Pixmap(image); page=output.new_page(width=pix.width,height=pix.height); page.insert_image(page.rect,stream=image)
        for b in page_data.blocks:
            if b.bbox and b.original_text:
                box=fitz.Rect(b.bbox.x1*pix.width,b.bbox.y1*pix.height,b.bbox.x2*pix.width,b.bbox.y2*pix.height)
                page.insert_textbox(box,b.normalized_text or b.original_text,fontsize=max(4,box.height*.7),render_mode=3)
    result=output.tobytes(); output.close(); return result
EXPORTERS={"txt":text,"md":markdown,"html":html_export,"docx":docx,"json":json_export}
MIMES={"txt":"text/plain; charset=utf-8","md":"text/markdown; charset=utf-8","html":"text/html; charset=utf-8","docx":"application/vnd.openxmlformats-officedocument.wordprocessingml.document","json":"application/json"}
