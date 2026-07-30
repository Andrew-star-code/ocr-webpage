import fitz
from docx import Document as WordDocument
from app.schemas.recognition import *
from app.services.exporters.base import ExportOptions
from app.services.exporters.registry import ExporterRegistry
from app.services.recognition.pipeline import RecognitionResult
from app.services.rendering.pages import RenderedPage
from PIL import Image
import io
def result(merged=False):
 cells=[TableCell(row_index=0,column_index=0,column_span=2 if merged else 1,is_header=True,text="Имя | №"),TableCell(row_index=0,column_index=1,is_header=True,text="Сумма")] if not merged else [TableCell(row_index=0,column_index=0,column_span=2,is_header=True,text="Имя | №")]
 table=TableBlock(id="t",reading_order=2,row_count=1,column_count=2,header_rows=1,rows=[TableRow(cells=cells)],bbox=BoundingBox(x1=.1,y1=.3,x2=.9,y2=.6))
 page=PageRecognition(page_number=1,width=200,height=100,blocks=[HeadingBlock(id="h",reading_order=1,original_text="Заголовок",bbox=BoundingBox(x1=.1,y1=.1,x2=.9,y2=.2)),table]);doc=DocumentRecognition(document_id="d",pages=[page],original_text="Заголовок",metadata=ProcessingMetadata(backend="ollama",model_profile="default",model_name="m"));image=io.BytesIO();Image.new("RGB",(200,100),"white").save(image,"PNG");return RecognitionResult(doc,[image.getvalue()],[RenderedPage(1,image.getvalue(),200,100,300,0,200,100)])
def test_registry_and_docx_merged_cells():
 registry=ExporterRegistry();assert registry.get("searchable_pdf")
 data=registry.get("docx").export(result(True),ExportOptions());word=WordDocument(io.BytesIO(data));assert len(word.tables)==1 and word.tables[0].cell(0,0).text=="Имя | №"
def test_markdown_merged_fallback_and_html_escape():
 registry=ExporterRegistry();markdown=registry.get("md").export(result(True),ExportOptions()).decode();assert "<table>" in markdown and "colspan=\"2\"" in markdown
 html=registry.get("html").export(result(),ExportOptions()).decode();assert "Имя | №" in html and "<thead>" in html
 escaped=result();escaped.document.pages[0].blocks[0].original_text="<script>";assert "<script>" not in registry.get("html").export(escaped,ExportOptions()).decode()
def test_searchable_pdf_preserves_image_and_cyrillic_layer():
 data=ExporterRegistry().get("searchable_pdf").export(result(),ExportOptions());pdf=fitz.open(stream=data,filetype="pdf");assert pdf.page_count==1 and pdf[0].get_images() and "Заголовок" in pdf[0].get_text()
def test_json_flags_remove_bbox_and_metadata():
 data=ExporterRegistry().get("json").export(result(),ExportOptions(include_bounding_boxes=False,include_processing_metadata=False)).decode();assert '"bbox"' not in data and '"metadata"' not in data
